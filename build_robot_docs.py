from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r"D:\vs-document\ROS2-Based Vision-Guided Autonomous Navigation Mobile Robot\output")
OUT.mkdir(exist_ok=True)

def font(run, size=10.5, bold=False, color=None, name='Calibri'):
    run.font.name = name; run.font.size = Pt(size); run.bold = bold
    rpr = run._element.get_or_add_rPr(); rf = rpr.rFonts
    rf.set(qn('w:ascii'), name); rf.set(qn('w:hAnsi'), name); rf.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if color: run.font.color.rgb = RGBColor(*color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); mar = tcPr.first_child_found_in('w:tcMar')
    if mar is None: mar=OxmlElement('w:tcMar'); tcPr.append(mar)
    for side,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=mar.find(qn('w:'+side))
        if node is None: node=OxmlElement('w:'+side); mar.append(node)
        node.set(qn('w:w'),str(val)); node.set(qn('w:type'),'dxa')

def fixed_table(doc, headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.autofit=False
    tblPr=t._tbl.tblPr; w=OxmlElement('w:tblW'); w.set(qn('w:w'),'9360'); w.set(qn('w:type'),'dxa'); tblPr.append(w)
    ind=OxmlElement('w:tblInd'); ind.set(qn('w:w'),'120'); ind.set(qn('w:type'),'dxa'); tblPr.append(ind)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.width=Inches(widths[i]); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; shade(c,'E8EEF5'); set_cell_margin(c)
        r=c.paragraphs[0].add_run(h); font(r,9.5,True,(31,77,120)); c.paragraphs[0].paragraph_format.space_after=Pt(0)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            c=cells[i]; c.width=Inches(widths[i]); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP; set_cell_margin(c)
            p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(1); p.paragraph_format.line_spacing=1.05; r=p.add_run(str(val)); font(r,8.8)
    return t

def add_page_field(p):
    r=p.add_run('第 '); font(r,9,False,(90,90,90)); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); p._p.append(fld); r=p.add_run(' 页'); font(r,9,False,(90,90,90))

def base_doc(running):
    d=Document(); sec=d.sections[0]; sec.top_margin=Inches(0.8); sec.bottom_margin=Inches(0.75); sec.left_margin=Inches(0.8); sec.right_margin=Inches(0.8); sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35)
    normal=d.styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); normal.font.size=Pt(10.5); normal.paragraph_format.space_after=Pt(5); normal.paragraph_format.line_spacing=1.22
    for name,size,color,before,after in [('Heading 1',16,(46,116,181),16,8),('Heading 2',13,(46,116,181),12,6),('Heading 3',11.5,(31,77,120),8,4)]:
        s=d.styles[name]; s.font.name='Calibri'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); s.font.size=Pt(size); s.font.color.rgb=RGBColor(*color); s.font.bold=True; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; rr=hp.add_run(running); font(rr,8.5,False,(105,105,105)); hp.paragraph_format.space_after=Pt(0)
    fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; add_page_field(fp)
    return d

def title(d, kicker, text, subtitle):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(68); p.paragraph_format.space_after=Pt(10); r=p.add_run(kicker); font(r,11,True,(46,116,181))
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(9); r=p.add_run(text); font(r,25,True,(11,37,69))
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(26); r=p.add_run(subtitle); font(r,12,False,(85,85,85))
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('版本 1.0 | 2026 年 8 月 27 日 | 面向实机落地与持续升级'); font(r,9.5,False,(110,110,110))
    d.add_page_break()

def para(d, text, style=None, boldlead=None):
    p=d.add_paragraph(style=style); p.paragraph_format.line_spacing=1.22
    if boldlead and text.startswith(boldlead):
        r=p.add_run(boldlead); font(r,10.5,True,(31,77,120)); r=p.add_run(text[len(boldlead):]); font(r)
    else: r=p.add_run(text); font(r)
    return p

def bullets(d, items):
    for x in items:
        p=d.add_paragraph(style='List Bullet'); p.paragraph_format.left_indent=Inches(.25); p.paragraph_format.first_line_indent=Inches(-.16); p.paragraph_format.space_after=Pt(2); r=p.add_run(x); font(r,10)

def h(d, level, text): d.add_heading(text,level=level)

def build_plan():
    d=base_doc('ROS2 AprilTag 自主泊车机器人 | 技术实施方案')
    title(d,'TECHNICAL IMPLEMENTATION PLAN','基于 ROS2、AprilTag 与 STM32 的视觉自主泊车机器人','从可演示 V1 到导航与具身任务平台 V3 的完整工程方案')
    h(d,1,'0. 文档用途、范围与结论')
    para(d,'本方案面向一台两轮差速、带编码器的低成本移动机器人。它不是“识别到二维码就前进”的玩具演示，而是以相机标定、相对位姿估计、速度闭环、状态机、安全停车、日志回放和可量化测试为核心的机器人系统。V1 的唯一闭环任务是：机器人从给定起点发现固定 AprilTag 后，完成搜索、对齐、低速接近、精确停车和结果上报。所有功能必须在实机上重复运行，并保留原始 rosbag、配置和测试记录。')
    para(d,'系统采用 RK3568 或笔记本作为 Linux/ROS2 上位机，STM32F103ZET6 作为实时底盘控制器。第一版不把 micro-ROS、激光 SLAM、Nav2、YOLO 和大模型同时塞入项目；这些能力全部作为接口兼容的后续版本。这样做能把 20 天内最易失败的机械、电源、驱动、通信和控制问题先解决，同时留下足够的技术上升空间。')
    fixed_table(d,['版本','核心能力','证明材料','不做什么'],[
        ['V1.0 泊车闭环','AprilTag 位姿、差速控制、停车状态机、安全保护','20 次测试表、视频、rosbag、参数文件','不依赖雷达、SLAM、微型 ROS 节点'],
        ['V1.5 工程化','URDF/TF、诊断、自动启动、配置管理、仿真回归','Gazebo 场景、故障注入记录、CI 构建','不夸大为“具身大模型”'],
        ['V2.0 自主导航','雷达、建图、AMCL/SLAM、Nav2、目标搜寻任务','地图、导航成功率、重规划记录','不把默认参数当成果'],
        ['V3.0 具身任务','视觉目标、任务语义、行为树、机械臂或配送扩展','结构化任务、反馈恢复、端到端演示','不直接由语言模型控制 PWM']
    ],[1.1,1.8,2.2,1.4])
    h(d,1,'1. 项目定位与验收边界')
    h(d,2,'1.1 场景定义')
    para(d,'标准演示场景为室内平整地面，入口处设置起始线，目标泊车位设置一张尺寸已知的 AprilTag。机器人接收“泊车到 Tag 0”任务后，先在限定角速度内原地扫描；检测到目标并连续确认后，以视觉相对位姿为依据控制车体转向和线速度；进入近距离区域后使用更低速度和更严格的姿态阈值；满足距离、横向偏差和航向误差条件后停车、锁定任务结果并上传诊断数据。若相机、通信、编码器或电池状态异常，必须无条件进入安全停车。')
    h(d,2,'1.2 V1 强制验收指标')
    bullets(d,['底盘能稳定执行直行、倒退、原地旋转和曲线运动；在同一电池电压区间内，低速指令不出现持续爬行或单轮失控。','上位机连续 300 ms 未刷新有效速度帧时，STM32 将左右电机目标速度置零并上报通信超时；急停输入触发后 PWM 立即关闭。','相机完成内参标定，标定文件随代码版本保存；Tag 的实际黑边尺寸、打印比例、安装高度和相机安装位姿均被记录。','在固定测试场地连续完成不少于 20 次泊车。建议门槛：任务成功率不低于 85%，最终距离误差不大于 8 cm，航向误差不大于 8 度；失败必须可从日志分类。','交付物包含源码、接线图、URDF、launch 文件、参数 YAML、协议说明、测试 CSV、至少一段无剪辑演示视频和一份复现实验说明。'])
    h(d,2,'1.3 诚实的能力边界')
    para(d,'V1 的决策属于规则控制和有限状态机，因此简历应称为“视觉感知驱动的自主泊车机器人”或“视觉定位与运动控制闭环”，而不是“具身大模型”或“VLA”。项目含金量来自于实机闭环、清晰的系统分层、故障处理和数据化评估；只有在后续加入开放词汇目标、任务解析、反馈重规划或策略学习并完成评测后，才适合进一步讨论具身智能。')
    h(d,1,'2. 参考项目与可借鉴边界')
    para(d,'linorobot2 提供了从底盘、机器人描述、仿真到 Nav2、SLAM Toolbox 与 robot_localization 的完整参考结构。项目应借鉴其“真实机器人与仿真复用同一描述和参数”的思路，不应直接复制其硬件假设或要求第一版就使用 micro-ROS。AprilTag ROS 2 节点可直接订阅校正图像和 CameraInfo，并发布检测数组及 TF；本项目应围绕该标准输出编写自己的 parking_controller，而不是改动检测器内部逻辑。')
    para(d,'Nav2 官方建模路径强调 URDF/TF、里程计、传感器、地图和导航插件的依赖顺序；这正是 V2 的升级路线。ros2_control 的 Resource Manager 与 Hardware Component 模型则适合在 V1 稳定后将串口底盘驱动包装为标准硬件接口，以获得 diff_drive_controller、joint_state_broadcaster 等标准控制器。原文中 Articulated Robotics 的 articubot_one 链接当前不可访问，因此不把其代码作为依赖；只保留“先建模、再仿真、再实机”的工程方法。')
    h(d,1,'3. 总体架构与数据流')
    para(d,'上位机负责非实时但计算量大的感知、状态机、任务管理、记录和可视化；STM32 负责固定周期的编码器采样、电机速度闭环、急停与通信看门狗。上位机只能下发“期望车体速度”和配置命令，绝不直接下发裸 PWM。STM32 永远拥有最后的安全决策权：任何上层崩溃、串口拔出或非法帧都应使底盘进入受控停车。')
    para(d,'数据路径：UVC 摄像头 -> camera_driver -> image_proc 校正 -> apriltag_ros -> detections/TF -> parking_controller -> /cmd_vel -> base_driver -> UART/CAN -> STM32；反向路径为编码器、轮速、电池、故障码 -> STM32 状态帧 -> base_driver -> /odom、/base_status、/battery_state、诊断日志。所有会影响决策的消息必须带时间戳，控制器对过期视觉结果、跳变位姿和重复帧执行拒绝或降级。')
    fixed_table(d,['层级','组件','职责','周期/约束'],[
        ['执行层','STM32 固件','编码器、PID、PWM、急停、CRC、看门狗','控制 100 Hz；通信超时 300 ms'],
        ['硬件抽象','base_driver','串口协议、里程计、状态发布、命令限幅','50 Hz 状态；20-50 Hz 指令'],
        ['感知层','camera/image_proc/apriltag_ros','采集、内参、畸变校正、Tag 位姿','15-30 FPS；图像与 CameraInfo 同时间基'],
        ['决策层','parking_controller','确认、搜索、对齐、接近、停车、恢复','20 Hz；状态显式可观测'],
        ['工程层','bringup/diagnostics/rosbag','启动编排、参数、健康检查、录制','异常不阻塞安全停车']
    ],[.9,1.35,3.35,.91])
    h(d,1,'4. 坐标系、机器人模型与运动学')
    para(d,'必须使用 REP 103/105 风格的右手坐标：x 前、y 左、z 上。最小 TF 树为 odom -> base_link -> camera_link -> camera_optical_frame；Tag 检测器可发布 camera_optical_frame -> tag36h11:0。V1 允许 odom 漂移，因为最终泊车主要依赖视觉相对位姿；但 odom 到 base_link 必须连续发布，便于 RViz、rosbag 和 V2 导航升级。不要把相机安装偏移硬编码在 C++ 中，应写入 xacro/URDF。')
    para(d,'差速底盘轮距 B、左右轮半径 r_l/r_r、编码器每电机转脉冲 PPR 和减速比均是标定参数。车体线速度与角速度映射为 v_l=v-wB/2、v_r=v+wB/2；反算时 v=(v_r+v_l)/2，w=(v_r-v_l)/B。实机中左右轮半径与死区并不完全相同，故应加入每轮比例系数、最小有效 PWM、加速度限制和方向反转保护，避免标称参数在低速下失效。')
    h(d,1,'5. ROS2 软件包与接口契约')
    para(d,'工作区建议采用一个元仓库：src/robot_description、robot_interfaces、base_driver、parking_controller、robot_bringup、robot_simulation、robot_tools。包之间只能通过消息、Action、参数和 TF 通信；不允许 parking_controller 直接打开串口或读取摄像头设备。每个包应有 README、launch、config、test 和最小示例，参数文件按 hardware、simulation、lab 三套环境分离。')
    fixed_table(d,['接口','方向','类型','约束'],[
        ['/cmd_vel','控制器 -> 底盘','geometry_msgs/Twist 或 TwistStamped','含最大速度、加速度和时间戳检查'],
        ['/odom','底盘 -> ROS2','nav_msgs/Odometry','odom -> base_link 与消息同源'],
        ['/camera/image_rect','相机 -> 检测器','sensor_msgs/Image','必须与 camera_info 同步'],
        ['/tag_detections','检测器 -> 控制器','AprilTagDetectionArray','使用 id、decision_margin、时间戳'],
        ['/base_status','底盘 -> 上层','自定义 BaseStatus','轮速、电压、故障、超时、序号'],
        ['/park','任务 -> 控制器','自定义 ParkToTag Action','反馈状态、取消、超时、最终误差']
    ],[1.35,1.1,1.85,3.1])
    para(d,'ParkToTag Action 建议目标字段为 tag_id、desired_distance_m、lateral_tolerance_m、yaw_tolerance_rad、timeout_s；反馈字段为 state、tag_visible、relative_x、relative_y、relative_yaw、command_v、command_w、retry_count；结果字段为 success、failure_code、final_distance、final_lateral_error、final_yaw_error、elapsed_time。Action 而不是单一 Topic 的价值在于任务可取消、可追踪、可写入测试结果，也便于 V2 由行为树调用。')
    h(d,1,'6. STM32 固件设计')
    para(d,'F103ZET6 资源充足，建议裸机定时器中断加轻量任务调度，或使用 FreeRTOS，但不要为了“用了 RTOS”制造不必要的竞态。推荐任务：1 kHz 编码器/保护采样、100 Hz 电机控制、50 Hz 通信解析与状态封包、10 Hz 电池检测与诊断。电机控制的关键变量用单一所有者维护；通信层只写目标速度邮箱，控制层在周期边界读取，避免串口 ISR 直接修改 PID 状态。')
    para(d,'速度 PID 采用位置式或增量式均可，但必须包含输出限幅、积分限幅、零速泄放、死区补偿、设定值斜坡和异常编码器检测。调参顺序先 P 后 D 最后少量 I；先让单轮无负载和落地低速稳定，再联调双轮直行。电池电压下降会改变电机特性，记录每次测试的电压，必要时加入简单前馈或低压限速。')
    fixed_table(d,['帧字段','字节建议','说明'],[
        ['SOF','2','固定 0xAA 0x55，便于帧同步'],['version','1','协议版本，升级时保持向后兼容'],['command/seq/length','1/1/1','命令字、序号、负载长度'],['payload','0-48','速度、参数或状态；小端、定点或明确 IEEE754'],['CRC16','2','覆盖 version 至 payload；错误帧丢弃'],['EOF','可选 1','仅用于调试，不替代长度和 CRC']
    ],[1.6,1.25,4.55])
    para(d,'命令至少包括 SET_TWIST、E_STOP、CLEAR_FAULT、SET_PID、PING；状态至少包括左右编码器增量、左右轮速、目标轮速、PWM、电池电压、故障位、急停位、最后命令年龄。对速度命令加速度限制和物理范围校验，CRC 错误计数、序号跳变和缓冲区溢出均写入诊断。串口默认 115200 可起步，若图像外的状态数据增加再提升到 460800；不要为了高波特率牺牲线缆、地线和抗干扰。')
    h(d,1,'7. 视觉定位与泊车控制')
    para(d,'相机标定使用 ROS image_pipeline 的 camera_calibration 或等价棋盘格流程。镜头焦距、分辨率、自动对焦状态与相机安装位置一旦变化，都必须重新标定。AprilTag 的 size 参数必须是实际黑色方块边长的米制值，不是纸张边长；推荐 36h11 家族、Tag 边长 10 至 16 cm、哑光打印并贴在硬质平板。先保存原始图像、矫正图像、CameraInfo 和检测结果，排除标定误差后再调控制器。')
    para(d,'检测确认不应只看一帧。建议连续 N=5 帧看到同一 id、decision_margin 高于阈值、hamming 不超过允许值、相对位姿的中位数变化合理时才进入对齐；若连续 M 帧丢失则停车并回到搜索。控制应使用 Tag 在 base_link 下的相对位置，先将 camera 到 base 的固定外参通过 TF 变换，再计算 x、y 与 yaw，绝不直接把像素中心偏差当作距离。')
    para(d,'搜索阶段只原地低速旋转，设置最大扫描时间；对齐阶段 v=0，以 w=clamp(k_yaw*yaw+k_lat*atan2(y,x),-w_max,w_max) 控制；接近阶段使用 v=clamp(k_x*(x-d_target),0,v_max)，同时叠加角速度修正。进入精对准区后降低 v_max 和 w_max，使用滞回阈值防止状态来回跳动。若目标位于相机后方、位姿不可信或控制命令超时，立即输出零速度。')
    h(d,1,'8. 状态机、安全与诊断')
    fixed_table(d,['状态','进入条件','动作','退出/失败'],[
        ['IDLE','上电完成','零速度、检查传感器与底盘','收到 Action 目标'],['SEARCH','任务开始且无可靠目标','低速旋转、统计扫描时间','连续确认或搜索超时'],['ALIGN','目标可靠且距离较远','原地对齐，v=0','偏角达标或目标丢失'],['APPROACH','姿态满足接近条件','限速前进并持续修正','到位、丢失、距离异常'],['FINE_DOCK','进入近距离阈值','低速精调、使用滞回','误差满足或超时'],['COMPLETE','泊车成功','发送零速度、固化结果','新任务或复位'],['FAULT','任何安全条件触发','PWM 安全停止、锁存故障','人工 clear_fault']
    ],[1.15,1.65,2.65,1.95])
    para(d,'安全机制分两层。STM32 层实施急停、通信超时、驱动器过流/欠压输入、编码器异常和 PWM 关断；ROS2 层实施相机失联、检测过期、控制器异常、任务超时和速度限幅。任何 ROS2 安全逻辑都不能替代 STM32 看门狗。物理急停建议串联在电机驱动器使能或电机电源路径，且按下后软件无法自行恢复。')
    h(d,1,'9. 标定、测试与数据闭环')
    para(d,'标定顺序：测量轮径和轮距 -> 单轮方向与 PPR 校验 -> 直线 2 m 试验校正左右轮比例 -> 原地转 360 度校正轮距 -> 相机内参 -> 相机到 base_link 外参 -> Tag 尺寸和安装位置 -> 泊车控制参数。每一步写入 calibration.yaml，并附日期、环境、工具和版本号。没有标定记录的“调通”不可复现。')
    fixed_table(d,['测试项','方法','通过准则','记录'],[
        ['通信超时','停止发送速度帧','300 ms 内轮速归零','状态帧、示波器或视频'],['急停','运动中按急停','PWM 关闭且故障锁存','故障码、恢复步骤'],['Tag 可见性','不同距离/角度/光照','定义可用工作区','检测率、margin、帧率'],['泊车重复性','固定起点 20 次','成功率和误差达到阈值','CSV、rosbag、无剪辑视频'],['异常恢复','遮挡 Tag/拔相机','停住并给出失败原因','状态迁移和日志'],['长期运行','连续运行 1 小时','无死锁、失控或资源泄漏','CPU、温度、错误计数']
    ],[1.25,2.15,1.9,2.1])
    para(d,'建议每个任务生成一条 run_id，保存 git commit、参数哈希、硬件版本、电压、起点、Tag id、最终误差、结果和失败码。rosbag2 只录制必要话题：/tf、/odom、/cmd_vel、/base_status、/tag_detections、/camera_info 和低帧率图像；同时保存控制器内部 CSV。这样既能回放定位问题，也不会因为录制全分辨率图像造成磁盘或 CPU 压力。')
    h(d,1,'10. 20 天实施计划与里程碑')
    fixed_table(d,['天数','目标','产出','风险闸门'],[
        ['1-2','硬件盘点、供电确认、仓库与协议','接线表、串口回环、URDF 骨架','未确认 RK3568 电源前不用电池供电'],['3-5','单轮/双轮编码器与 PID','轮速曲线、超时停车','低速不稳先暂停上层开发'],['6-7','base_driver、odom、TF、RViz','/cmd_vel 到实机闭环','时间戳和方向必须正确'],['8-10','相机标定、AprilTag、控制状态机','首次自主泊车','先固定光照与固定 Tag'],['11-12','安全、诊断、异常恢复','故障注入表','所有故障默认停车'],['13-15','重复性测试、参数固化、仿真','20 次数据与 Gazebo 场景','失败先分类再调参'],['16-18','V1.5：自动启动、日志、README','可复现部署','不新增高风险硬件'],['19-20','演示、复盘、简历材料','视频、图、测试报告','只写实测数据']
    ],[.85,1.75,2.4,2.4])
    h(d,1,'11. V1.0 至 V3.0 升级路线')
    para(d,'V1.5 先补工程化：为底盘与相机写 URDF/xacro，建立 Gazebo Harmonic 模型，使用同一份接口和参数文件做仿真回归；为 launch 加入生命周期与诊断；将串口驱动拆为 transport、protocol、base_state 三层。V2 再加入 2D 雷达、IMU、robot_localization、SLAM Toolbox、AMCL 与 Nav2。雷达与相机位置只通过 TF/URDF 描述，parking_controller 改为 Nav2 到“搜索区”后的子任务，仍使用同一个 ParkToTag Action 完成最终泊车。')
    para(d,'V2.5 可将 base_driver 封装为 ros2_control SystemInterface，向外暴露左右轮 velocity command 与 position/velocity state，接入 diff_drive_controller；此时仍保留现有二进制协议和 STM32 PID，不做底层推倒重来。V3 再引入 YOLO/RKNN 目标识别、结构化 JSON 任务、BehaviorTree.CPP 编排、视觉验证和失败重规划。语言模型最多生成经过 schema 校验的高层任务，不得绕过限速、可达性与安全状态机直接输出底盘控制。')
    h(d,1,'12. 风险清单与决策原则')
    bullets(d,['最常见失败不是 ROS2，而是电机供电、共地、编码器接线、驱动器额定电流和轮胎打滑。先用电流表和低占空比确认硬件，再进入软件联调。','RK3568 的镜像、Ubuntu 版本、GPU/NPU 驱动和 USB 带宽存在板卡差异；V1 可以先在笔记本运行 ROS2，等链路稳定后再迁移，避免把开发阻塞在系统镜像上。','不购买激光雷达也能完成 V1；若为了“看起来高级”提前加雷达，会同时引入驱动、TF、地图、定位和 Nav2 调参，严重挤压泊车闭环与测试时间。','所有简历指标必须由测试 CSV 产生。没有实际测量时写“完成 20 次重复测试并记录指标”，不要编造成功率、延迟或定位精度。'])
    h(d,1,'13. 详细实施细则与排障手册')
    details = '''
13.1 环境基线。推荐把笔记本或 RK3568 的系统、ROS 发行版、内核版本、相机驱动版本、串口设备名和 udev 规则写入 environment.md。首次验证时使用 Ubuntu 24.04 arm64/x86_64 与 ROS2 Jazzy；若 RK3568 的厂商镜像不是标准 Ubuntu，先在笔记本完成全部 ROS2 包开发和仿真，再将可独立编译的工作区迁移到板端。不要把厂商 SDK、RKNN、相机驱动和机器人业务代码混在一个目录。每次部署前以脚本检查 ROS_DOMAIN_ID、串口权限、相机设备、磁盘空间和时间同步，避免“换一台机器就不能跑”。
13.2 仓库与版本管理。根目录放置 README、docs、src、config、launch、scripts、hardware、data 和 test_results；hardware 保存接线图、器件数据表和协议版本，data 不提交大体积原始 rosbag，只保存样例和下载说明。每个可运行版本建立 Git tag，例如 v1.0-pid、v1.1-odom、v1.2-tag-parking。参数文件必须由 launch 显式加载，禁止依赖个人 bashrc 的隐式设置。提交信息说明硬件变化、参数变化或行为变化，测试 CSV 中保存 commit id，这样视频、代码与指标可以一一对应。
13.3 电源与地线。电机电源从电池经保险丝、急停和主开关进入 H 桥；RK3568、STM32、相机和电机不得直接串在同一根细杜邦线上。逻辑电源由独立稳压模块提供，电机电源和逻辑地只在规划的公共地连接。电机驱动器附近增加合适的电解电容，编码器和串口线尽量远离电机线；如果发生相机掉线、串口乱码或板端重启，先用万用表检查电压跌落和地线，而不是先怀疑 ROS2。所有带电插拔仅在主电源断开时进行。
13.4 编码器验证。上电后先不接轮子，手动转动单轮并在串口打印左右计数和方向；确认正向定义与 ROS base_link 的 x 正方向一致。再给每个电机固定 PWM，记录 10 秒脉冲数、轮速和电流，建立左右电机特性表。若两轮同样 PWM 差异过大，不要通过“给右轮更大速度命令”掩盖，应检查机械阻力、轮胎、减速箱和编码器信号。速度估计的采样窗过短会抖动，过长会滞后；推荐以 10 ms 或 20 ms 原始增量配合一阶滤波，并在低速区使用计数累积窗口。
13.5 PID 调试纪律。先关闭 I、D，只调整 P 让轮速对阶跃命令有响应但不持续振荡；再小幅加入 D 抑制超调，最后只加入足以消除稳态误差的 I。调参时固定电池状态、载荷和地面，单次只改一个参数。输出必须经限幅和斜坡，尤其是从正转切换反转时先经过零速保持区。出现尖锐噪声、轮胎打滑、H 桥过热或电机堵转时立即停止，控制器参数不能补偿机械故障。将最终 PID、采样周期、PWM 频率和死区作为版本化参数保存。
13.6 里程计质量。V1 不以全局里程计精度作为泊车依据，但 odom 仍是后续 Nav2 的地基。直线测试采用 2 m 和 5 m 两个距离，分别测量前进与后退误差；旋转测试做多次 360 度原地转向，校正轮距而不是随意修改角速度系数。左右轮半径、轮距和编码器计数的任何修正都应写入 calibration.yaml。发布 Odometry 时 pose 使用 odom 坐标系、twist 使用 base_link 坐标系，并保证 odom 到 base_link 的 TF 与消息时间戳一致；这是 Nav2、RViz 和机器人定位模块的基本约定。
13.7 相机与 AprilTag 实操。相机固定在刚性支架上，镜头尽量朝前并略向下，避免车体振动造成松动。标定采样需覆盖画面中心、四角和不同倾斜角，采集结束后检查重投影误差和畸变系数是否异常。将一张 Tag 贴到平整硬板上，避免纸张卷曲；在不同距离、左右偏角、光照和运动速度下测量 detection rate 与 decision_margin，画出工作范围。若检测坐标方向反了，先检查 camera_optical_frame 与外参，而不要在控制公式中人为加负号。视觉结果的中位数滤波窗口应小于控制动态响应，不能以大窗口制造“看起来稳定”的滞后。
13.8 控制器防抖与限速。停车控制常见问题是远处可见但近处 Tag 被相机视野裁掉，或目标在车体中心附近时横向误差受噪声影响反复翻转。解决方法是分区限速、目标丢失滞回、到位后锁存 COMPLETE，以及在近距离时把图像/Tag 可靠性作为必要条件。所有控制输出经过 v_max、w_max、线加速度、角加速度和最小可动速度限制；当指令低于电机死区时，应明确输出零速度而不是发送无效 PWM。状态机中每个状态都设置最大停留时间，任何超时都有确定失败码和零速度动作。
13.9 故障码设计。建议按位划分：bit0 通信超时，bit1 急停，bit2 电池欠压，bit3 左编码器异常，bit4 右编码器异常，bit5 H 桥告警，bit6 参数非法，bit7 固件看门狗复位。上位机收到故障后只负责显示、录制和取消任务，不应循环下发 clear_fault。清故障需要人工确认急停释放、通信恢复和电压正常，STM32 再接受带序号的 CLEAR_FAULT。诊断界面至少显示最后命令年龄、帧错误计数、PID 输出、轮速和电压，使故障能在视频中被观察。
13.10 仿真策略。Gazebo 模型不追求逼真的外观，而应复用真实 URDF 的轮距、轮径、相机和未来雷达位置。第一阶段在仿真中验证 TF 树、Action 取消、状态机超时和速度限幅；第二阶段人为注入 Tag 丢失、传感器延迟和通信断开，验证安全状态。仿真通过不是实机通过，但它能防止每次改代码都直接冒险跑车。真实硬件与仿真使用相同的话题名称和参数接口，只在 hardware.yaml、simulation.yaml 中切换传输和传感器插件。
13.11 自动启动与可维护性。稳定后使用 systemd 管理 bringup，服务先等待串口和相机设备出现，再启动 ROS2 launch；失败后有限次数重启并保留 journal。不要在 systemd 服务中写死 /dev/ttyUSB0，应使用 udev 规则按 VID/PID 或序列号创建 /dev/robot_base。日志目录按日期和 run_id 分层，加入磁盘容量阈值，防止 rosbag 写满系统盘。提供 stop_robot.sh，先取消任务、发布零速度、等待底盘确认，再停止进程；不能依赖直接拔电源作为正常关机。
13.12 评测报告模板。每次正式测试写明测试目的、场地尺寸、地面材质、光照、Tag 尺寸、起点、软件版本、硬件版本和电池电压。结果表至少记录成功/失败、失败码、耗时、最终距离、横向误差、航向误差、平均帧率、最大命令年龄和人工备注。报告区分开发试跑与正式统计，不能把调整参数时选出的最好一次作为平均结果。成功率给出分母，例如“17/20 次，85%”，并列出三次失败分别是目标丢失、通信超时还是姿态未收敛。
13.13 简历与答辩素材。项目完成后准备四类证据：一张分层架构图，一段从启动到泊车成功的无剪辑视频，一张 20 次测试的误差/成功率图，以及一页故障恢复记录。简历职责可写成“设计 ROS2 上位机与 STM32 下位机分层架构，完成编码器闭环、CRC 串口协议、AprilTag 相对位姿泊车状态机与超时急停；通过多次实车测试量化泊车成功率和误差”。只有数字已经测量时才写具体数值；答辩中主动说明 V1 不含 SLAM、V2 的升级接口，反而更显得工程判断可靠。
13.14 V2 导航的接入顺序。购买雷达后仍先完成静态 TF、/scan 可视化和时间戳检查，再做手动建图，之后才考虑 SLAM Toolbox、地图保存、AMCL 和 Nav2。导航到搜索区与 Tag 最终泊车是两个不同层级：Nav2 负责全局/局部路径与避障，ParkToTag Action 负责末端视觉对接。通过行为树或任务管理器在 Nav2 成功后调用泊车 Action，并把泊车失败回传为可重试或人工接管。不要让 Nav2 与 parking_controller 同时抢占 /cmd_vel；使用速度仲裁或明确的控制权状态。
13.15 从小车到具身平台。后续接入 YOLO/RKNN 时，先把检测结果转为标准化目标对象，包括类别、置信度、空间位置、时间戳、可达性和稳定性，再让任务管理器选择目标。自然语言只转换为经过 schema 校验的 action、object、destination 等字段；每一步均由视觉确认和状态机约束。若未来增加机械臂，可复用任务接口、诊断、日志、数据集格式和失败恢复框架，小车项目就不再是孤立 demo，而成为移动操作系统的一部分。'''.strip().split('\n')
    for item in details:
        para(d,item)
    h(d,1,'14. 参考资料与使用说明')
    refs=['ROS 2 Jazzy Ubuntu 安装文档：https://docs.ros.org/en/jazzy/Installation/Alternatives/Ubuntu-Install-Binary.html','AprilTag ROS 2 Node：https://github.com/christianrauch/apriltag_ros','linorobot2：https://github.com/linorobot/linorobot2','Nav2 First-Time Robot Setup Guide：https://docs.nav2.org/rolling/configuration_and_development/first_time_robot_setup_guide/','ros2_control Jazzy 文档：https://control.ros.org/jazzy/doc/getting_started/getting_started.html','ROS image_pipeline / camera_calibration：https://docs.ros.org/en/jazzy/p/image_pipeline/']
    bullets(d,refs)
    para(d,'访问日期：2026-08-27。上述开源项目用于理解标准接口、目录组织、仿真和导航集成方式；本项目的硬件接线、通信协议、控制状态机、参数和测试记录应自行设计并在仓库中完整保留。')
    return d

def build_purchase():
    d=base_doc('机器人项目硬件购买方案 | 小车与桌面机械臂')
    title(d,'HARDWARE PROCUREMENT GUIDE','机器人项目硬件购买方案','小车 V1/V2 与第二个桌面机械臂项目的低成本采购、兼容性与预算')
    h(d,1,'1. 先盘点已有设备，再下单')
    para(d,'已知可复用设备包括：野火开发板（型号待核实）、STM32F103ZET6 开发板、若干 STM32 最小系统板，以及粤嵌机构的 RK3568 开发板。F103ZET6 足够承担小车底盘控制；不要因为“有野火板”就额外购买 MCU。RK3568 可承担最终上位机，但在未确认系统镜像、供电规格、USB 主机口数量和 Ubuntu/ROS2 兼容性前，建议先用笔记本搭 ROS2 环境。')
    fixed_table(d,['现有物','建议用途','必须确认','不建议立即购买'],[
        ['STM32F103ZET6','小车电机、编码器、串口、安全控制','供电电压、定时器引脚、下载器、板载 3.3V 能力','重复 MCU'],['野火开发板','备用实验、传感器验证或机械臂控制','具体型号、PWM/串口资源、资料包','未确认型号前的扩展盾'],['STM32 最小板','协议/驱动单元测试','晶振、下载、IO 电平','作为主控的冗余替代'],['粤嵌 RK3568','ROS2 上位机、相机、日志、后续 RKNN','输入电压、散热、USB、系统镜像','盲买电源或转接板']
    ],[1.35,1.85,2.15,1.55])
    h(d,1,'2. 小车 V1 必买清单：推荐档')
    para(d,'推荐档的目标是让 PID、编码器和低速泊车稳定，不是追求最低单价。核心原则是选择带 AB 相霍尔编码器的金属减速直流电机，并根据堵转电流选择驱动器。若选普通 TT 电机配 TB6612FNG，成本最低但轮速一致性和耐久性较差；若选 JGA25-370 一类金属电机，必须避免把高堵转电流直接压给 TB6612。')
    fixed_table(d,['类别','推荐规格/关键词','数量','参考价格（元）','理由与注意'],[
        ['底盘','2WD 金属底盘，含轮毂与万向轮','1','60-120','优先平整刚性底板，预留上层板和相机支架'],['电机','JGA25-370 12V，AB 霍尔编码器，成对同型号','2','100-180/对','比 TT 更适合速度闭环；确认减速比、PPR、轴径'],['驱动','BTS7960 x2 或双路大电流 H 桥','1 套','40-90','按堵转电流留至少 1.5 倍余量；3.3V 逻辑兼容'],['主控','现有 STM32F103ZET6','1','0','负责控制与安全，不另购'],['相机','UVC USB 720p/1080p 固定焦或可锁焦','1','45-100','不要用频繁自动对焦的摄像头；需 UVC 免驱'],['串口','CP2102/CH340 USB-TTL，3.3V 可选','1','8-20','RK3568 调试和烧录备用；TX/RX 交叉且共地'],['电池','成品 3S 11.1V 锂电池包+BMS+充电器','1','80-160','避免裸 18650 拼装；确认持续放电电流'],['降压','12V 转 5V 5A 同步降压模块','1','15-35','仅在确认 RK3568 输入需求后给上位机供电'],['安全','急停蘑菇头、保险丝座+保险丝、电源开关','1 套','20-45','急停切断电机使能/电源，软件不能绕过'],['线材','端子、硅胶线、杜邦线、热缩管、螺丝','1 套','40-80','电机线与信号线分开走线，统一接地']
    ],[.85,2.35,.5,.9,1.9])
    h(d,2,'2.1 小车 V1 预算与下单顺序')
    para(d,'不含已有 MCU 和 RK3568 的推荐档合计约 410 至 830 元，取决于电机、底盘、电池和相机质量。若预算只有 250 至 400 元，可用带编码器 TT 电机、TB6612FNG 和轻量底盘完成验证，但应把它定位为“原型档”，并预留更换驱动器/电机的安装孔位。下单顺序应为：底盘电机驱动与电源 -> 相机和安全件 -> 线材与支架；雷达、IMU、RGB-D 相机均不是 V1 的必买项。')
    h(d,2,'2.2 必须避免的采购错误')
    bullets(d,['购买“无编码器电机”后再试图依赖视觉做速度闭环；没有轮速反馈，停车重复性和里程计升级都会很差。','只看电机额定电流，不看堵转电流，导致 H 桥发热、掉压或烧毁。','用开发板 5V 或 USB 口给电机/大舵机供电；电机电源、逻辑电源必须分开稳压但共地。','未确认 RK3568 是 5V、9V 还是 12V 输入就接电池。第一阶段使用原厂适配器，确认手册后再接入车载供电。','购买自动对焦低端摄像头后用于标定；焦距变化会使内参失效，泊车表现漂移。'])
    h(d,1,'3. 小车升级件：只在 V1 稳定后采购')
    fixed_table(d,['升级目标','器件','预算（元）','采购条件','带来的能力'],[
        ['姿态融合','MPU6050 或 BMI088','10-120','先有稳定编码器与 TF','角速度辅助、EKF 实验'],['Nav2/建图','LD06/LD19 或 RPLIDAR、YDLIDAR 2D 雷达','300-900','先完成 URDF/odom；确认 Linux 驱动','SLAM、AMCL、避障'],['电源可靠性','电压/电流采样、DC-DC、接线端子、保险','60-180','开始长期运行测试','低压保护、可维护性'],['计算性能','散热片/风扇、稳定 USB Hub、SSD','80-300','RK3568 负载和供电测试通过','日志、视觉推理、稳定 USB'],['标准控制接口','CAN 收发器或隔离串口','15-100','线长/干扰明显或多控制器','总线扩展与可靠通信']
    ],[1.2,2.2,1.05,1.8,1.25])
    para(d,'linorobot2 的当前文档列举了 LD06、LD19、RPLIDAR 和 YDLIDAR 等常见 2D 雷达支持路径，但“驱动可用”不等于“实机可导航”。购买雷达后还需要完成安装高度、TF、时间戳、噪声、轮廓 footprint、局部代价地图和速度参数调优。因此它应是 V2 的明确里程碑，不是第一个视频前必须具备的装饰。')
    h(d,1,'4. 第二项目：桌面机械臂硬件方案')
    para(d,'第二项目建议做“视觉识别与桌面抓取/分拣”，与小车共享笔记本或 RK3568、USB 相机、ROS2 工作区、日志与任务状态机。为了 20 天内可落地，第一版选择轻载、4 至 5 自由度、固定桌面工位的舵机机械臂；它可证明视觉到动作闭环，但不应被包装成高精度工业机械臂。')
    fixed_table(d,['类别','经济型推荐','数量','参考价格（元）','说明'],[
        ['机械臂本体','4DOF/5DOF 金属或加厚亚克力套件','1','120-260','优先有固定底座和夹爪；套件舵机仅作起点'],['关节舵机','肩/肘用 DS3218 或同级金属齿；轻关节 MG90S','4-6','120-260','不要让 SG90 承担大臂负载'],['PWM 扩展','PCA9685 16 路 PWM','1','8-25','STM32 做轨迹插补，PCA9685 做脉冲输出'],['独立电源','6V 10A 成品开关电源或电池降压','1','50-110','舵机与 MCU 逻辑共地，电源独立'],['视觉','复用小车 UVC 相机 + 俯视支架','1','20-60','相机固定后做手眼/平面标定'],['末端反馈','微动开关或 FSR 薄膜压力传感器','1-2','8-35','辅助判断夹爪闭合/抓取'],['工作物','彩色积木、哑光盒子、棋盘格板','1 套','30-70','控制物体尺寸和反光，先做可重复任务'],['安全','限位、急停、线束固定件','1 套','25-60','机械臂同样要有物理断电与软件软限位']
    ],[.85,2.35,.5,.9,1.9])
    h(d,2,'4.1 机械臂的两条升级路线')
    para(d,'路线 A 是低成本舵机臂：用 STM32、PCA9685、USB 相机和固定工位，完成颜色/类别识别、平面坐标映射、预抓取/抓取/放置、视觉复检和失败重试。这是第二项目的推荐起点。路线 B 是 LeRobot SO-101 或同类智能总线舵机臂：其 follower 使用 6 个 STS3215 总线舵机，能提供更好的关节反馈与遥操作/数据采集基础，但成本、打印件和调试工作量都明显更高；应在 V1 分拣闭环完成后再考虑。')
    h(d,1,'5. 采购前兼容性检查单')
    bullets(d,['查清 RK3568 板卡型号、原厂输入电压、电源接口极性、系统版本、可用 USB Host 数量与是否能运行 Ubuntu 24.04/ROS2 Jazzy；截图或拍照留档。','电机卖家页面必须能确认电压、减速比、编码器 AB 相电平、每圈脉冲、轴径、空载电流和堵转电流；成对购买同批次。','H 桥需要同时满足电机堵转电流、供电电压和 3.3V PWM/方向信号。若逻辑高电平不明确，准备电平转换或选明确支持 3.3V 的模块。','电池必须有成品保护板、匹配充电器、可承受两台电机启动电流的持续放电能力；加保险丝并使用阻燃收纳。','相机需确认 UVC、Linux 可识别、固定焦距/可锁焦、分辨率和帧率；收到后第一时间用 v4l2-ctl 测试。'])
    h(d,1,'6. 最终采购建议')
    para(d,'第一笔只买小车 V1 推荐档，且雷达延后。若预算允许，优先把钱从“炫目的传感器”投入到带编码器金属电机、合适的 H 桥、独立电源、保险与急停、固定焦相机和可靠线材上。这些物件决定你能否获得稳定、可测的实机数据。第二笔再购买机械臂本体、关节舵机与独立 6V 电源；相机、上位机和部分安全材料可复用。')
    para(d,'购买后先完成一页硬件资产表：型号、卖家链接、到货日期、实测电压、引脚、驱动版本和照片。该表会成为接线图、README、故障定位和简历答辩的证据。所有价格均为 2026 年 8 月常见渠道参考区间，不替代下单当天的实际价格与规格确认。')
    h(d,1,'7. 参考资料')
    bullets(d,['linorobot2 硬件与传感器参考：https://github.com/linorobot/linorobot2','ROS2 控制硬件抽象：https://control.ros.org/jazzy/doc/getting_started/getting_started.html','LeRobot SO-101 官方装配与电机说明：https://huggingface.co/docs/lerobot/so101','MoveIt 2 官方入门资料：https://moveit.picknik.ai/main/doc/tutorials/getting_started/getting_started.html'])
    return d

if __name__=='__main__':
    p=build_plan(); p.save(OUT/'ROS2_AprilTag_STM32_视觉自主泊车机器人_完整项目方案.docx')
    q=build_purchase(); q.save(OUT/'机器人项目硬件购买方案_小车与桌面机械臂.docx')
